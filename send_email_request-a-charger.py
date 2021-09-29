#Send a email and image as an attachment with python 
#importing the modules
import os
#send an email
import smtplib, ssl
from email.message import EmailMessage, Message
from email.mime.text import MIMEText
from email.mime.image import MIMEImage
from email.mime.multipart import MIMEMultipart
#edit word docs
from docx import Document #pip install python-docx
from docx.shared import Pt #used for styling the font size
#doc to pdf converter
from docx2pdf import convert #pip install docx2pdf
#pdf to img
from pdf2image import convert_from_path #pip install pdf2image


#creating an email class
#order to run class: edit templeate --> convert docx to pdf to img --> send img as email
class Email():
    def __init__(self):
        #content for email: sender, receiver, password, message, and subject
        self.SENDER_EMAIL_ADDRESS = 'requset-a-charger@computer4u.com' 
        self.RECEIVER_EMAIL_ADDRESS = ''
        self.PASSWORD = 'SC59UB8LY'

        self.MESSAGE = """Hi, 
        You recently requested to rent a charger for a device.
        Click the photo to view the instructions"""
        self.SUBJECT  = 'Verification Code'

        #content for file
        self.TEMPLEATE = 'Templeate.docx' #file name for the Templeate docx
        self.NEW_FILE_NAME = 'VerificationCode' #the new name for the file once its has been edited
        self.fontSize = 28 #font size for verification code
        self.defaultFontSize = 12 #font size for deafult text

        self.VERIFICATION_CODE = ''

    #edit the templeate with the verification code
    def EditTempleate(self):
        doc = Document(self.TEMPLEATE)
        #looking for name in the templeate to replace it
        for p in doc.paragraphs:
            if 'name' in p.text:
                p.text = '' #clearing the txt
                p.add_run('Hello {},'.format(self.RECEIVER_EMAIL_ADDRESS)).font.size = Pt(self.defaultFontSize) #setting the text with a font size

        #looking for verification code in the templeate to replace it
        for p in doc.paragraphs:
            if 'Your Verification Code' in p.text:
                p.text = '' #clearing the text
                p.add_run(self.VERIFICATION_CODE).font.size = Pt(self.fontSize) #setting the text with a font size
        doc.save(self.NEW_FILE_NAME+'.docx')

    #covert the doc to pdf then to img
    def ConvertDoc(self):
        try:
            try:
                convert(self.NEW_FILE_NAME+'.docx') #convert docx to pdf, only works if microsoft word is installed
            except:
                os.system('abiword --to=pdf filetoconvert.doc') #if we are running this on a raspberry pi, then use abiword. abiword needs to be installed though
        except:
            raise Exception('Microsoft Word or Abiword is not installed')
        #convert pdf to jpg
        images = convert_from_path(self.NEW_FILE_NAME+'.pdf')
        for i in range(len(images)):
            images[i].save(self.NEW_FILE_NAME+'.jpg', 'JPEG') 
    
    #function for sending the email
    def sendEmail(self):
        #opening the image file to send
        with open(self.NEW_FILE_NAME+'.jpg', 'rb') as f:
            img_data = f.read()

        #setting up an email message
        msg = MIMEMultipart()

        #Subject, from, and to (information for email)
        msg['Subject'] = self.SUBJECT
        msg['From'] = self.SENDER_EMAIL_ADDRESS
        msg['To'] = self.RECEIVER_EMAIL_ADDRESS

        #attaching the text an image to the email   
        text = MIMEText(self.MESSAGE)
        msg.attach(text)
        image = MIMEImage(img_data, name=os.path.basename(self.NEW_FILE_NAME+'.jpg'))
        msg.attach(image)
        
        #connecting to server and sending email
        server = smtplib.SMTP('smtp.mail.com', 587) #logging into an outlook server (this is only for the sender. Meaning the sender has to be outlook. You can send an email to any domain)
        server.starttls()
        server.login(self.SENDER_EMAIL_ADDRESS, self.PASSWORD) #login to the email
        print('login successful')
        server.send_message(msg) #send message
        print('email has been sent to ', self.RECEIVER_EMAIL_ADDRESS)
        server.quit() #quitting the sever
        print('server quitting')


email = Email()
email.VERIFICATION_CODE = '93812fe'
email.RECEIVER_EMAIL_ADDRESS = 'epresent@emeryweiner.org'
email.EditTempleate()
email.ConvertDoc()
email.sendEmail()