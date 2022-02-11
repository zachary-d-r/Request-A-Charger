import os
import random

import smtplib, ssl
from email.message import EmailMessage, Message
from email.mime.text import MIMEText
from email.mime.image import MIMEImage
from email.mime.multipart import MIMEMultipart

from docx import Document #pip install python-docx
from docx.shared import Pt #used for styling the font size
from docx2pdf import convert #pip install docx2pdf
from pdf2image import convert_from_path #pip install pdf2image

class verification:
    
    MESSAGE = ''  # This is your message
    SUBJECT = 'Verification Number'  # This is your subject

    codeNums = ['1', '2', '3', '4', '5', '6', '7', '8', '9', 'A', 'B', 'C', 'F', 'L', 'Y', 'E', 'P', 'X', 'U', 'W', 'R', 'G', 'S', 'V', 'Z']  # Characters that can be in a verification code
    verificationCode = ''  # The code
    toAddress = ''  # Who we are sending the email to
    fromAddress = ''  # Who is sending the email
    #password = 'SC59UB8LY'  # Password to login to smtp TODO remove
    verified = False
    isVerified = False

    # Initialize the object by setting the correct email addresses
    def __init__(self, to, fromEmail, password):
        self.toAddress = to
        self.fromAddress = fromEmail
        self.password = password
        self.verified = False

        self.MESSAGE = """Hi, 
        You recently requested to rent a charger for a device.
        Click the photo to view the instructions.

        Sincerely,
        The Upper School Robotics Team"""
        self.SUBJECT  = 'Verification Code: {code}'

        # Content for file
        self.TEMPLATE = 'Verification-Files/Template.docx' # File name for the Template docx
        self.NEW_FILE_NAME = 'Verification-Files/VerificationCode' # The new name for the file once its has been edited
        self.fontSize = 28 # Font size for verification code
        self.defaultFontSize = 12 # Font size for deafult text


    # Generate a random verification code
    def getVerificationCode(self):
        self.verificationCode = ''  # Reset the verification code

        # Add a random character from the list of characters above
        for _ in range(5):
            self.verificationCode += self.codeNums[random.randint(0, len(self.codeNums)-1)]

    
    # Edit the template with the verification code
    def EditTemplate(self):
        doc = Document(self.TEMPLATE)
        # Looking for name in the template to replace it
        for p in doc.paragraphs:
            if 'name' in p.text:
                p.text = '' # Clearing the txt
                p.add_run('Hello {},'.format(self.toAddress)).font.size = Pt(self.defaultFontSize) # Setting the text with a font size
        
        # Looking for verification code in the template to replace it
        for p in doc.paragraphs:
            if 'Your Verification Code' in p.text:
                p.text = '' # Clearing the text
                p.add_run(self.verificationCode).font.size = Pt(self.fontSize) # Setting the text with a font size
        doc.save(self.NEW_FILE_NAME+'.docx')

    
    # Covert the doc to pdf then to img
    def ConvertDoc(self):
        try:
            try:
                convert(self.NEW_FILE_NAME+'.docx') #convert docx to pdf, only works if microsoft word is installed
            except:
                os.system('abiword --to=pdf filetoconvert.doc') #if we are running this on a raspberry pi, then use abiword. abiword needs to be installed though
        except:
            raise Exception('Microsoft Word or Abiword is not installed')

        #convert pdf to jpg

        # Try to convert using poppler from root
        try:
            images = convert_from_path(self.NEW_FILE_NAME+'.pdf')

        # Catch an exception if poppler was not found in root, and use the downloaded version
        except:
            images = convert_from_path(self.NEW_FILE_NAME+'.pdf', poppler_path=r'poppler-0.68.0\bin') #C:\Users\zacha\OneDrive\Desktop\poppler-0.68.0_x86\   #Change to work with computer using: https://stackoverflow.com/questions/18381713/how-to-install-poppler-on-windows #Don't need to specify location for linux
        
        # Save the image
        finally:
            for i in images:
                i.save(self.NEW_FILE_NAME+'.jpg', 'JPEG')


    # Email the verification code to the client
    def sendEmail(self):

        """

        # Opening the image file to send
        with open(self.NEW_FILE_NAME+'.jpg', 'rb') as f:
            img_data = f.read()
        """
        msg = MIMEMultipart()  # Setting up an email message
        

        #TODO remove
        """
        # If the user did not want to enter the password in the function parameters then they can enter it as user input
        if self.password == None:
            self.password = input('Please enter your password\n')
        """


        # Subject, from, and to (information for email)
        msg['Subject'] = self.SUBJECT.format(code=self.verificationCode)
        msg['From'] = self.fromAddress
        msg['To'] = self.toAddress

        # Attaching the text and image to the email
        text = MIMEText(self.MESSAGE)
        msg.attach(text)
        #image = MIMEImage(img_data, name=os.path.basename(self.NEW_FILE_NAME+'.jpg'))
        #msg.attach(image)

        # Logging into an outlook server (this is only for the sender. Meaning the sender has to be mail.com. You can send an email to any domain)
        with smtplib.SMTP('smtp.gmail.com', 587) as server:
            server.ehlo()   # Fixed the problem
            server.starttls()
            server.login(self.fromAddress, self.password) # Login to the email
            server.send_message(msg) # Send message


    def sendVerificationCode(self):
        #self.EditTemplate()
        #self.ConvertDoc()
        self.sendEmail()